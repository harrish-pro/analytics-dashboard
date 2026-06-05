from docx import Document


def create_docx_report(
    report_data,
    file_path
):

    doc = Document()

    # REPORT TITLE

    doc.add_heading(
        "AI Analytics Report",
        level=1
    )

    # DATASET OVERVIEW

    profile = report_data["profile"]

    doc.add_heading(
        "Dataset Overview",
        level=2
    )

    doc.add_paragraph(
        f"""
Rows: {profile['rows']}
Columns: {profile['columns']}
Missing Values: {profile['missing_values']}
Duplicates: {profile['duplicates']}
"""
    )

    # EXECUTIVE SUMMARY

    doc.add_heading(
        "Executive Summary",
        level=2
    )

    doc.add_paragraph(
        report_data["summary"]
    )

    # HEALTH SCORE

    health_score = report_data["health_score"]

    doc.add_heading(
        "Dataset Health Score",
        level=2
    )

    doc.add_paragraph(
        f"{health_score}%"
    )

    # AI INSIGHTS

    insights = report_data["insights"]

    doc.add_heading(
        "AI Business Insights",
        level=2
    )

    for insight in insights:

        doc.add_paragraph(
            insight
        )

    # RECOMMENDATIONS

    recommendations = report_data["recommendations"]

    doc.add_heading(
        "Recommendations",
        level=2
    )

    for recommendation in recommendations:

        doc.add_paragraph(
            recommendation
        )

    # SAVE FILE

    doc.save(file_path)